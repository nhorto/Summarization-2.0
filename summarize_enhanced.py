"""
Enhanced Multi-Document Summarization Script (v3)
-------------------------------------------------
This script creates professional weekly consulting summary reports from transcript files.

IMPROVEMENTS OVER V2:
- Enhanced prompts for better narrative flow and professional tone
- Automatic Word document (.docx) generation with formatting
- Opening paragraph generation (thank you + week focus)
- Closing paragraph generation (summary + availability)
- Better section organization matching consultant report style

PIPELINE:
1. Summarizes each transcript in transcripts/ folder into daily summaries
2. Synthesizes all daily summaries into a master weekly report
3. Generates formatted Word document in output/ folder

ANTI-HALLUCINATION SAFEGUARDS:
- Only uses information explicitly present in source text
- No assumptions, guesses, or invented content
- Critical instructions repeated in every LLM call
"""

import os
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent

TRANSCRIPTS_DIR = BASE_DIR / "transcripts"
PROCESSED_TRANSCRIPTS_DIR = BASE_DIR / "transcripts_processed"
DAILY_SUMMARIES_DIR = BASE_DIR / "summaries_daily"
MASTER_SUMMARY_DIR = BASE_DIR / "summaries_master"
OUTPUT_DIR = BASE_DIR / "output"

# Model & API
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL_NAME = os.getenv("OPENAI_MODEL_NAME", "gpt-5-nano-2025-08-07")

# Chunking parameters
CHUNK_CHAR_LENGTH = 15000
CHUNK_CHAR_OVERLAP = 800


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def ensure_dirs() -> None:
    DAILY_SUMMARIES_DIR.mkdir(parents=True, exist_ok=True)
    MASTER_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_TRANSCRIPTS_DIR.mkdir(parents=True, exist_ok=True)


def load_llm() -> ChatOpenAI:
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not set in the environment.")
    return ChatOpenAI(model=MODEL_NAME)


def read_srt(path: Path) -> str:
    """Extract text from .srt file, skipping index and timestamp lines."""
    lines: List[str] = []
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if not line or line.isdigit() or "-->" in line:
                continue
            lines.append(line)
    return "\n".join(lines)


def read_vtt(path: Path) -> str:
    """Extract text from .vtt (WebVTT) file, skipping headers and timestamps."""
    lines: List[str] = []
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if (
                not line
                or line.startswith("WEBVTT")
                or line.startswith("NOTE")
                or line.startswith("STYLE")
                or line.startswith("REGION")
                or "-->" in line
                or line.isdigit()
            ):
                continue
            lines.append(line)
    return "\n".join(lines)


def read_txt(path: Path) -> str:
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_transcript(path: Path) -> str:
    if path.suffix.lower() == ".srt":
        return read_srt(path)
    if path.suffix.lower() == ".vtt":
        return read_vtt(path)
    return read_txt(path)


def process_transcripts() -> List[Path]:
    """
    Convert raw .srt/.vtt files into clean daily .txt files in PROCESSED_TRANSCRIPTS_DIR.
    Files sharing the same base filename (stem) are combined into one daily transcript.
    """
    if not TRANSCRIPTS_DIR.exists():
        print(f"⚠️  Transcripts directory does not exist: {TRANSCRIPTS_DIR}")
        return []

    raw_files = sorted(
        [p for p in TRANSCRIPTS_DIR.iterdir() if p.suffix.lower() in {".srt", ".vtt"}]
    )

    if not raw_files:
        print("⚠️  No .srt or .vtt transcript files found to process.")
        return []

    # Start fresh each run so processed folder tracks the current raw inputs
    for existing in PROCESSED_TRANSCRIPTS_DIR.glob("*.txt"):
        try:
            existing.unlink()
        except OSError as exc:
            print(f"   ⚠️  Could not remove previous processed file {existing.name}: {exc}")

    grouped_files = defaultdict(list)
    for file_path in raw_files:
        grouped_files[file_path.stem].append(file_path)

    processed_paths: List[Path] = []

    print("\nStep 0: Normalizing transcripts (.srt/.vtt -> daily .txt)\n")
    for stem in sorted(grouped_files.keys()):
        day_files = sorted(grouped_files[stem])
        combined_sections: List[str] = []

        for file_path in day_files:
            try:
                text = load_transcript(file_path).strip()
            except Exception as exc:
                print(f"   ❌ Error reading {file_path.name}: {exc}")
                continue

            if not text:
                print(f"   ⚠️  {file_path.name} contained no usable text, skipping.")
                continue

            combined_sections.append(text)

        if not combined_sections:
            print(f"   ⚠️  No clean transcript content for {stem}, skipping output file.")
            continue

        combined_text = "\n\n".join(combined_sections)
        out_path = PROCESSED_TRANSCRIPTS_DIR / f"{stem}.txt"
        try:
            with out_path.open("w", encoding="utf-8") as f:
                f.write(combined_text)
            processed_paths.append(out_path)
            print(f"   ✅ {out_path.name} created from {[p.name for p in day_files]}")
        except Exception as exc:
            print(f"   ❌ Could not write processed transcript {out_path.name}: {exc}")

    if not processed_paths:
        print("⚠️  No processed transcripts were created.")

    return processed_paths


def chunk_text(text: str, length: int = CHUNK_CHAR_LENGTH, overlap: int = CHUNK_CHAR_OVERLAP) -> List[str]:
    """Simple character-based chunking with overlap."""
    text = text.strip()
    if not text:
        return []

    chunks: List[str] = []
    start = 0
    n = len(text)

    while start < n:
        end = min(start + length, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)

    return chunks


# ---------------------------------------------------------------------------
# Enhanced Prompts
# ---------------------------------------------------------------------------

DAILY_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant creating a topic-organized consulting summary from a single day's transcript.

CRITICAL INSTRUCTIONS (MANDATORY - NO EXCEPTIONS):
- Only use information that is EXPLICITLY stated in the transcript
- DO NOT make assumptions, guesses, or draw conclusions not directly supported by transcript content
- DO NOT add information that was not discussed or mentioned
- If something is unclear or not mentioned, DO NOT include it
- Maintain STRICT factual accuracy - fabricating details is unacceptable

STYLE INSTRUCTIONS:
- Organize the summary into functional consulting topics, such as:
  • Administrative Information / Administration
  • Estimating (Accessories, Paint & Cleaning, Estimate Summary, Labor Codes, etc.)
  • Project Management (Trimble Connect, Schedules, Drawing Log, RFIs, Contracts, etc.)
  • Purchasing
  • Production Control (Import Files, Work Packages, Routes, Cut Lists, etc.)
  • Inspections / QA/QC
  • PowerFab Go
  • Detailing Standards
  • IT Security / Data Integration
  • Any other relevant topics

- Use professional consultant tone and structure:
  • Brief opening statement (1-2 sentences) describing the focus of this session
  • Topic headers in Title Case or ALL CAPS (e.g., "Estimating" or "PROJECT MANAGEMENT")
  • Under each topic: 3-10 detailed narrative bullet points describing:
    - What was reviewed, discussed, or demonstrated
    - Key findings, observations, or determinations
    - Configurations made or changes implemented
    - Decisions reached or recommendations provided
    - Any explicit next steps mentioned
  • Optional brief closing statement summarizing progress

- DO NOT create global sections titled:
  "What We Looked At and Covered During the Week"
  "What We Found and Determined Needed to Happen"
  "What We Accomplished and Did"
  "What You Need to Do Moving Forward"

  These are conceptual frameworks, not section headings. Instead, naturally weave this information into topic-specific bullets.

CONTENT RULES:
- Group all content by consulting topic
- Within each topic's bullets, naturally include:
  • What was reviewed/covered
  • What was found/determined
  • What was accomplished/configured
  • Next steps (only if explicitly mentioned)
- Do NOT invent action items or decisions not clearly present in the transcript
- Use professional, clear, and concise language
- Write in past tense (what was done/discussed)
"""

DAILY_USER_PROMPT_TEMPLATE = """You will receive raw transcript text from a Tekla PowerFab consulting session between a consultant and client.

Create a professional topic-based consulting summary following the STYLE and CRITICAL INSTRUCTIONS in the system prompt.

Transcript:
\"\"\"{transcript_chunk}\"\"\""""

MASTER_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant synthesizing multiple daily consulting summaries into one comprehensive weekly consulting report.

CRITICAL INSTRUCTIONS (MANDATORY - NO EXCEPTIONS):
- Only use information EXPLICITLY stated in the provided daily summaries
- DO NOT make assumptions, guesses, or draw conclusions not directly supported by the summaries
- DO NOT add information that was not discussed or mentioned in the summaries
- Maintain STRICT factual accuracy - fabricating details is unacceptable
- Do NOT invent recommendations, next steps, or action items not present in the source material

STYLE INSTRUCTIONS:
- Match the professional structure and tone of a Tekla PowerFab weekly consulting report:

  • DO NOT include an opening paragraph yet - that will be generated separately

  • Topic-based sections organized by functional area:
    - Administrative Information / Administration
    - Estimating (with subsections: Accessories, Paint & Cleaning, Labor Codes, etc.)
    - Project Management (with subsections: Trimble Connect, Schedules, RFIs, etc.)
    - Purchasing
    - Production Control (with subsections: Import Files, Routes, Cut Lists, etc.)
    - Inspections / QA/QC
    - PowerFab Go
    - Any other relevant topics from the summaries

  • Under each topic or subtopic, use detailed narrative bullet points that:
    - Describe what was reviewed, discussed, or demonstrated
    - Explain key findings, observations, or determinations
    - Detail configurations made or changes implemented
    - Note decisions reached or recommendations provided
    - Include explicit next steps or action items (only if mentioned in summaries)

  • DO NOT include a closing paragraph yet - that will be generated separately

- DO NOT use global section headings like:
  "What We Looked At and Covered During the Week"
  "What We Found and Determined Needed to Happen"
  "What We Accomplished and Did"
  "What You Need to Do Moving Forward"

  These concepts should be naturally woven into the topic-specific bullets.

CONTENT RULES:
- Consolidate and de-duplicate information across all daily summaries
- Group related items under the most appropriate topic header
- Use subtopic headers when a topic has multiple distinct areas (e.g., "Estimating" with "Accessories", "Labor Codes", etc.)
- Within each topic's bullets, naturally structure content as:
  • What was reviewed/covered
  • What was found/determined needed to happen
  • What was accomplished/configured/trained on
  • What needs to happen moving forward (only when explicitly stated)
- Do NOT invent new topics, tasks, or recommendations
- Preserve all important details while eliminating redundancy
- Use professional, clear, and precise language
- Write in past tense for completed actions
- Use present tense only for current state descriptions
"""

MASTER_USER_PROMPT_TEMPLATE = """You will receive multiple daily consulting summaries from the same client engagement (one week of consulting work).

Using ONLY those summaries, synthesize a comprehensive weekly consulting report that follows the STYLE and CRITICAL INSTRUCTIONS from the system prompt.

Remember: Do NOT include opening or closing paragraphs - focus only on the topic-based content sections.

Daily summaries:
\"\"\"{daily_summaries_text}\"\"\""""

OPENING_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant writing the opening paragraph for a weekly consulting report.

CRITICAL INSTRUCTIONS:
- Only use information EXPLICITLY stated in the provided report content
- DO NOT invent client names, specific dates, or details not present in the content
- Maintain professional, warm, and appreciative tone
- Keep it concise (2-5 sentences)

STYLE:
- Start with gratitude (thanking the client for the opportunity)
- Briefly state the focus/scope of the week's work
- List key areas covered (based on the topics in the content)
- Set a positive, professional tone for the rest of the report

EXAMPLE STRUCTURE:
"This summary report outlines the work completed during my on-site visit with your team. Throughout the week, we evaluated your current use of Tekla PowerFab, identified key areas for improvement, and implemented changes across [list main topics]. The following sections detail the discussions we had, the adjustments we made, and the recommendations provided for continued progress."
"""

OPENING_USER_PROMPT_TEMPLATE = """Based on the following weekly consulting report content, write a professional opening paragraph that thanks the client and summarizes the week's focus.

Report content:
\"\"\"{report_content}\"\"\""""

CLOSING_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant writing the closing paragraph for a weekly consulting report.

CRITICAL INSTRUCTIONS:
- Only use themes and topics EXPLICITLY present in the provided report content
- DO NOT invent new recommendations or action items not mentioned in the report
- Maintain warm, professional, and supportive tone
- Keep it concise (3-6 sentences)

STYLE:
- Thank the client for their engagement and openness
- Reinforce the key theme: consistent system usage across the organization
- Emphasize the importance of deciding how to use the system and communicating that standard
- Offer continued support and availability
- End on a positive, encouraging note

EXAMPLE STRUCTURE:
"Thank you for allowing me to work with your team this week. I want to emphasize what I believe is the most important takeaway: decide how you want to use the system, determine which areas will be adopted, and clearly communicate that expectation to everyone. [Additional reinforcement of key themes from the report]. I truly enjoyed my time with your team and appreciate your willingness to pursue meaningful improvement. Please feel free to reach out if you need any assistance as you implement these changes or if you would like me to return in the future."
"""

CLOSING_USER_PROMPT_TEMPLATE = """Based on the following weekly consulting report content, write a professional closing paragraph that thanks the client, reinforces key themes, and offers continued support.

Report content:
\"\"\"{report_content}\"\"\""""


# ---------------------------------------------------------------------------
# LLM Call Helpers
# ---------------------------------------------------------------------------

def call_llm(llm: ChatOpenAI, system_prompt: str, user_prompt: str) -> str:
    resp = llm.invoke([
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ])
    return resp.content if hasattr(resp, "content") else str(resp)


def summarize_transcript_file(path: Path, llm: ChatOpenAI) -> str:
    """Create a topic-based daily summary for a single transcript file."""
    raw = load_transcript(path)
    chunks = chunk_text(raw)
    if not chunks:
        return ""

    partial_summaries: List[str] = []

    for i, chunk in enumerate(chunks, start=1):
        print(f"    - Summarizing chunk {i}/{len(chunks)} for {path.name}...")
        user_prompt = DAILY_USER_PROMPT_TEMPLATE.format(transcript_chunk=chunk)
        chunk_summary = call_llm(llm, DAILY_SYSTEM_PROMPT, user_prompt)
        partial_summaries.append(chunk_summary.strip())

    if len(partial_summaries) == 1:
        return partial_summaries[0]

    # Compress multiple chunk summaries into one
    merged_text = "\n\n---\n\n".join(partial_summaries)
    print(f"    - Creating compressed daily summary for {path.name}...")
    user_prompt = DAILY_USER_PROMPT_TEMPLATE.format(transcript_chunk=merged_text)
    final_summary = call_llm(llm, DAILY_SYSTEM_PROMPT, user_prompt)
    return final_summary.strip()


def create_master_summary(llm: ChatOpenAI, daily_summary_paths: List[Path]) -> str:
    """Create a weekly master summary from multiple daily summary files."""
    texts: List[str] = []
    for p in daily_summary_paths:
        with p.open("r", encoding="utf-8", errors="ignore") as f:
            texts.append(f"=== {p.name} ===\n{f.read().strip()}")

    combined = "\n\n\n".join(texts)
    user_prompt = MASTER_USER_PROMPT_TEMPLATE.format(daily_summaries_text=combined)
    master_summary = call_llm(llm, MASTER_SYSTEM_PROMPT, user_prompt)
    return master_summary.strip()


def generate_opening_paragraph(llm: ChatOpenAI, report_content: str) -> str:
    """Generate professional opening paragraph based on report content."""
    user_prompt = OPENING_USER_PROMPT_TEMPLATE.format(report_content=report_content[:3000])  # Use first 3000 chars
    return call_llm(llm, OPENING_SYSTEM_PROMPT, user_prompt).strip()


def generate_closing_paragraph(llm: ChatOpenAI, report_content: str) -> str:
    """Generate professional closing paragraph based on report content."""
    # Sample content from throughout the report for better thematic understanding
    sample_content = report_content[:2000] + "\n...\n" + report_content[-1000:]
    user_prompt = CLOSING_USER_PROMPT_TEMPLATE.format(report_content=sample_content)
    return call_llm(llm, CLOSING_SYSTEM_PROMPT, user_prompt).strip()


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------

def create_word_document(content: str, opening: str, closing: str, output_path: Path) -> None:
    """Generate formatted Word document from report content."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add date at the top
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)

    # Add opening paragraph
    doc.add_paragraph()  # Blank line
    opening_para = doc.add_paragraph(opening)
    opening_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add blank line before content
    doc.add_paragraph()

    # Process the main content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if this is a major section header (ALL CAPS or Title Case with no bullet)
        if line.isupper() or (line[0].isupper() and not line.startswith('•') and not line.startswith('-')):
            # Add some space before major headers (except the first one)
            if len(doc.paragraphs) > 3:
                doc.add_paragraph()

            heading = doc.add_paragraph(line)
            heading.style = 'Heading 1'
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(12)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 0, 0)

        # Check if this is a bullet point
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            bullet_text = line.lstrip('•-*').strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(6)

        # Regular paragraph
        else:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add closing paragraph
    doc.add_paragraph()  # Blank line
    closing_para = doc.add_paragraph(closing)
    closing_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add signature line
    doc.add_paragraph()
    signature = doc.add_paragraph("Sincerely,")
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save document
    doc.save(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    ensure_dirs()

    print("=" * 70)
    print(" Enhanced Tekla Consulting Summarization (v3)")
    print("=" * 70)
    print(f"Transcripts directory:      {TRANSCRIPTS_DIR}")
    print(f"Processed transcripts dir:  {PROCESSED_TRANSCRIPTS_DIR}")
    print(f"Daily summaries directory:  {DAILY_SUMMARIES_DIR}")
    print(f"Master summaries directory: {MASTER_SUMMARY_DIR}")
    print(f"Output directory:           {OUTPUT_DIR}")
    print()

    processed_transcripts = process_transcripts()
    if not processed_transcripts:
        print("❌ Unable to continue without processed daily transcripts.")
        return

    llm = load_llm()
    transcript_files = processed_transcripts

    # Step 1: Create daily summaries
    daily_summary_paths: List[Path] = []

    print("Step 1: Creating daily topic-based summaries\n")
    for path in transcript_files:
        print(f"▶ Processing transcript: {path.name}")
        try:
            summary_text = summarize_transcript_file(path, llm)
            if not summary_text:
                print(f"   ⚠️  Empty summary for {path.name}, skipping.")
                continue

            out_name = f"{path.stem}_summary.txt"
            out_path = DAILY_SUMMARIES_DIR / out_name
            with out_path.open("w", encoding="utf-8") as f:
                f.write(summary_text)
            print(f"   ✅ Saved daily summary -> {out_path}")
            daily_summary_paths.append(out_path)
        except Exception as e:
            print(f"   ❌ Error summarizing {path.name}: {e}")
        print()

    if not daily_summary_paths:
        print("⚠️  No daily summaries were created; skipping master summary.")
        return

    # Step 2: Create master summary
    print("\nStep 2: Creating weekly master summary from daily summaries\n")
    try:
        master_text = create_master_summary(llm, daily_summary_paths)
        master_out_path = MASTER_SUMMARY_DIR / "master_summary.txt"
        with master_out_path.open("w", encoding="utf-8") as f:
            f.write(master_text)
        print(f"✅ Master summary saved -> {master_out_path}")
    except Exception as e:
        print(f"❌ Error creating master summary: {e}")
        return

    # Step 3: Generate opening paragraph
    print("\nStep 3: Generating opening paragraph\n")
    try:
        opening_text = generate_opening_paragraph(llm, master_text)
        print(f"✅ Opening paragraph generated")
    except Exception as e:
        print(f"❌ Error generating opening: {e}")
        opening_text = ""

    # Step 4: Generate closing paragraph
    print("\nStep 4: Generating closing paragraph\n")
    try:
        closing_text = generate_closing_paragraph(llm, master_text)
        print(f"✅ Closing paragraph generated")
    except Exception as e:
        print(f"❌ Error generating closing: {e}")
        closing_text = ""

    # Step 5: Create Word document
    print("\nStep 5: Creating Word document\n")
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = OUTPUT_DIR / f"Weekly_Consulting_Summary_{timestamp}.docx"
        create_word_document(master_text, opening_text, closing_text, docx_path)
        print(f"✅ Word document created -> {docx_path}")
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")

    print("\n" + "=" * 70)
    print("All done! Your weekly consulting summary is ready.")
    print("=" * 70)


if __name__ == "__main__":
    main()
