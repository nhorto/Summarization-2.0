"""
Resume Summarization Script
---------------------------
This script picks up where summarize_enhanced.py left off.
It allows you to:
1. Specify which days still need daily summaries
2. Generate the master summary from all daily summaries
3. Create the final Word document

Usage:
    python summarize_resume.py

Configure the DAYS_TO_PROCESS list below to specify which days need summaries.
"""

import os
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from openai import OpenAI

# ---------------------------------------------------------------------------
# Configuration - EDIT THIS SECTION
# ---------------------------------------------------------------------------

# Specify which days still need daily summaries (leave empty if all daily summaries are done)
# Example: ["Tuesday", "Wednesday"]
DAYS_TO_PROCESS = ["Tuesday", "Wednesday"]

# Set to True to skip daily summaries entirely and just create master summary + Word doc
SKIP_DAILY_SUMMARIES = False

# ---------------------------------------------------------------------------
# Paths (same as main script)
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent

TRANSCRIPTS_DIR = BASE_DIR / "transcripts"
PROCESSED_TRANSCRIPTS_DIR = BASE_DIR / "transcripts_processed"
DAILY_SUMMARIES_DIR = BASE_DIR / "summaries_daily"
MASTER_SUMMARY_DIR = BASE_DIR / "summaries_master"
OUTPUT_DIR = BASE_DIR / "output"

# ---------------------------------------------------------------------------
# Model & API
# ---------------------------------------------------------------------------

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL_NAME = os.getenv("OPENAI_MODEL_NAME", "gpt-4o-mini")

CEREBRAS_API_KEY = os.getenv("CEREBRAS_API_KEY")
CEREBRAS_MODEL_NAME = os.getenv("CEREBRAS_MODEL_NAME", "gpt-oss-120b")
CEREBRAS_BASE_URL = "https://api.cerebras.ai/v1"

CHUNK_CHAR_LENGTH = 15000
CHUNK_CHAR_OVERLAP = 800

# ---------------------------------------------------------------------------
# Prompts (copied from main script)
# ---------------------------------------------------------------------------

DAILY_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant creating a detailed topic-organized summary from a single day's consulting session transcript.

CRITICAL INSTRUCTIONS (MANDATORY - NO EXCEPTIONS):
- Only use information that is EXPLICITLY stated in the transcript
- DO NOT make assumptions, guesses, or draw conclusions not directly supported by transcript content
- DO NOT add information that was not discussed or mentioned
- If something is unclear or not mentioned, DO NOT include it
- Maintain STRICT factual accuracy - fabricating details is unacceptable
- Preserve ALL technical details: numbers, values, tool names, density figures, specifications, etc.

ORGANIZATION INSTRUCTIONS:
- Identify the organic topics and themes that emerged during this session
- DO NOT force content into predefined categories - let the topics emerge naturally from the transcript
- Group all related discussions under the most appropriate topic heading
- Create as many topic sections as needed to comprehensively capture the session
- Use descriptive topic headers that reflect what was actually discussed (e.g., "Stainless Steel Conversion and Weight Calculations" not just "Estimating")

CONTENT INSTRUCTIONS:
- Be extremely detailed and comprehensive - these summaries feed the final report
- Capture what was reviewed, discussed, demonstrated, or analyzed
- Document key findings, observations, determinations, and concerns
- Record configurations made, changes implemented, or settings adjusted
- Note decisions reached, recommendations provided, or professional judgments made
- Include any explicit next steps, action items, or follow-up tasks mentioned
- Preserve technical details: specific numbers, density values, tool names, version numbers, etc.
- Include context about WHY things were discussed or WHY they matter
- Note any concerns, risks, limitations, or issues identified
- Document questions raised, uncertainties, or items requiring further investigation

STYLE INSTRUCTIONS:
- Use clear, professional consultant language
- Organize content by topic/theme with descriptive headers
- Use bullet points, sub-bullets, and paragraphs as needed for clarity
- Write in past tense for completed actions
- Use present tense for current state descriptions
- Include professional observations and analytical language when appropriate
- It's acceptable to use structured formats since these are internal working documents

Remember: The goal is comprehensive capture of all session content in an organized way that can be synthesized into the final weekly report. No detail is too small if it was explicitly discussed.
"""

DAILY_USER_PROMPT_TEMPLATE = """
You will receive raw transcript text from a Tekla PowerFab consulting session between a consultant and client.

Create a professional topic-based consulting summary following the STYLE and CRITICAL INSTRUCTIONS in the system prompt.

Transcript:
\"\"\"{transcript_chunk}\"\"\"

"""

MASTER_SYSTEM_PROMPT = """
You are a professional Tekla PowerFab consultant writing a comprehensive weekly consulting report. This report will be delivered to the client, so it must be polished, professional, and read like a thoughtfully written document, not a bullet-point checklist.

CRITICAL INSTRUCTIONS (MANDATORY - NO EXCEPTIONS):
- Only use information EXPLICITLY stated in the provided daily summaries
- DO NOT make assumptions, guesses, or draw conclusions not directly supported by the summaries
- DO NOT add information that was not discussed or mentioned in the summaries
- Maintain STRICT factual accuracy - fabricating details is unacceptable
- Preserve ALL technical details: numbers, values, tool names, specifications, etc.
- Do NOT invent recommendations, next steps, or action items not present in the source material

WRITING STYLE - THIS IS CRITICAL:
Your writing must be narrative, flowing prose - like a professional consultant writing a report by hand after careful thought. This is NOT a bullet-point document.

GOOD EXAMPLE (narrative prose):
"Shapes, Grades, and Sizes were reviewed as part of the administrative setup, with a focus on improving consistency and ensuring the database supports accurate estimating, purchasing, and inventory workflows. While progress was made during the week, this area is not yet fully complete and will require additional follow-up.

Duplicate grades are still present in the database and will need to be cleaned up. Some consolidation was discussed, but the cleanup process has not been fully completed. Until duplicate grades are removed, there is an increased risk of inconsistent material selection across estimates, purchase orders, and inventory transactions."

BAD EXAMPLE (bullet-point checklist):
"Shapes/Grades/Sizes
- What was reviewed / covered
  - Reviewed shapes, grades, and sizes as part of administrative setup
  - Focus on improving consistency for estimating, purchasing, and inventory
- Key findings / observations
  - Duplicate grades still present in database
  - Cleanup process not fully completed
- Next steps
  - Complete cleanup of duplicate grades"

STRUCTURE REQUIREMENTS:
- Write in flowing paragraphs with natural transitions between ideas
- Use topic section headers that reflect the organic themes from the week's work
- Each topic section should read like a story: what was done, why it matters, what was found, what the implications are, and what needs to happen next
- Use bullet points ONLY when listing specific items where bullets genuinely add clarity (rarely)
- Include professional judgment, analytical observations, and consultant perspective
- Discuss concerns, risks, limitations, and implications in narrative form
- Explain WHY things matter and what the consequences or benefits are

TOPIC ORGANIZATION:
- Identify the natural themes and topics that emerged across the week's daily summaries
- DO NOT force content into predefined categories like "Estimating, Project Management, Purchasing"
- Let the topics emerge organically from what was actually discussed
- Use descriptive, specific topic headers (e.g., "Duplicate grades are still present..." not just "Grades")
- Some topics may have subtopics - organize these naturally with secondary headers when needed
- Topics should be ordered logically, with related areas grouped together

NARRATIVE DEPTH:
- Provide context and background for why work was undertaken
- Explain findings and their implications or consequences
- Discuss concerns, risks, or limitations identified
- Note what is complete, what is in progress, and what remains to be done
- Include professional observations about data quality, consistency, or process gaps
- Describe not just WHAT was done, but WHY it matters and WHAT it means for the client

TECHNICAL PRECISION:
- Preserve all specific numbers, values, formulas, and calculations
- Include tool names, version numbers, and technical specifications
- Reference specific features, settings, or configurations by their correct names
- Maintain accurate technical terminology throughout

PROFESSIONAL TONE:
- Write as a consultant speaking to the client
- Use "we" when describing collaborative work
- Use past tense for completed work
- Use present tense for current state or ongoing conditions
- Include analytical language: "this indicates," "this suggests," "this will require," "it was observed that"
- Express professional concerns or recommendations with appropriate gravity
- Be direct but professional when noting issues or gaps

FORMAT REQUIREMENTS:
- DO NOT include an opening paragraph (generated separately)
- DO NOT include a closing paragraph (generated separately)
- Focus on the topic-based content sections only
- Use section headers formatted as plain text (e.g., "Administrative Information")
- Subtopic headers can be used when a major topic has distinct areas (e.g., under "Estimating": "Labor Standards", "Material Pricing")
- Write in paragraph form - this is the default, not the exception
- Use bullets only when listing specific items enhances readability (e.g., a list of tools, a series of rates)

Remember: This report represents your professional consulting work. It should read like a thoughtful, carefully written document that demonstrates expertise, attention to detail, and genuine consulting value. The client should feel they received a professional analysis, not just a transcription of notes.
"""

MASTER_USER_PROMPT_TEMPLATE = """
You will receive multiple daily consulting summaries from the same client engagement (one week of consulting work).

Using ONLY those summaries, synthesize a comprehensive weekly consulting report that follows the WRITING STYLE and STRUCTURE REQUIREMENTS from the system prompt.

Your goal is to produce a narrative, professional document that reads like a hand-written consultant report with flowing prose, not a bullet-point checklist.

Key requirements:
1. Identify organic topics and themes from the daily summaries
2. Write in narrative paragraphs with natural flow and transitions
3. Include professional judgment, analytical observations, and implications
4. Preserve all technical details and specifications
5. Use bullets sparingly - only when genuinely appropriate for listing items
6. DO NOT include opening or closing paragraphs (generated separately)
7. Provide depth and context - explain why things matter and what they mean

Daily summaries:
\"\"\"{daily_summaries_text}\"\"\"
"""

OPENING_SYSTEM_PROMPT = """You are a professional Tekla PowerFab consultant writing the opening paragraph for a weekly consulting report.

CRITICAL INSTRUCTIONS:
- Only use information EXPLICITLY stated in the provided report content
- DO NOT invent client names, specific dates, or details not present in the content
- Maintain professional, warm, and appreciative tone
- Keep it concise but substantive (2-5 sentences)

STYLE:
- Start with gratitude, thanking the client for the opportunity to work with their team
- Briefly describe the scope or focus of the week's work in narrative form
- Reference the main topics or themes covered during the week
- Set a positive, professional tone for the rest of the report
- Write as flowing prose, not a list

EXAMPLE STRUCTURE:
"Thank you for the opportunity to support your team this week. The focus covered [major theme 1], [major theme 2], and [major theme 3] within Tekla PowerFab, including [specific areas with a few concrete examples]. We [briefly describe the nature of the work - reviewed, configured, discussed, implemented, etc.] across [list main topics naturally], advancing your team's [capability/readiness/understanding] in these critical areas."

The opening should feel like a consultant speaking directly to the client, setting the stage for the detailed report that follows.
"""

OPENING_USER_PROMPT_TEMPLATE = """Based on the following weekly consulting report content, write a professional opening paragraph that thanks the client and summarizes the week's focus in narrative prose.

Use the topics and themes from the report to craft 2-5 sentences that introduce the work completed during the week.

Report content:
\"\"\"{report_content}\"\"\"
"""

CLOSING_SYSTEM_PROMPT = """
You are a professional Tekla PowerFab consultant writing the closing paragraph for a weekly consulting report.

CRITICAL INSTRUCTIONS:
- Only use themes and topics EXPLICITLY present in the provided report content
- DO NOT invent new recommendations or action items not mentioned in the report
- Maintain warm, professional, and supportive tone
- Keep it concise but meaningful (3-6 sentences)

STYLE:
- Thank the client for their engagement, collaboration, and openness
- Reinforce a key theme or insight from the week (based on report content)
- Emphasize the importance of consistency, communication, and adoption where relevant
- Offer continued support and availability in a genuine, professional manner
- End on a positive, encouraging note about the progress made and path forward
- Write as flowing narrative prose

EXAMPLE STRUCTURE:
"Thank you for allowing me to work with your team this week. I want to emphasize what I believe is the most important takeaway: [key insight or theme from the report, such as: decide how you want to use the system, establish standards, ensure consistent adoption, etc.]. [Additional 1-2 sentences reinforcing themes from the report or noting progress/momentum]. I truly enjoyed my time with your team and appreciate your willingness to pursue meaningful improvement. Please feel free to reach out if you need any assistance as you implement these changes or if you would like me to return in the future."

The closing should feel personal, professional, and should reinforce the value of the work completed while offering genuine support for next steps.
"""

CLOSING_USER_PROMPT_TEMPLATE = """Based on the following weekly consulting report content, write a professional closing paragraph that thanks the client, reinforces key themes from the week, and offers continued support.

Identify the major themes from the report and use those to craft a meaningful, personalized closing (3-6 sentences) in narrative prose.

Report content:
\"\"\"{report_content}\"\"\"
"""


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def ensure_dirs() -> None:
    DAILY_SUMMARIES_DIR.mkdir(parents=True, exist_ok=True)
    MASTER_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_TRANSCRIPTS_DIR.mkdir(parents=True, exist_ok=True)


def load_cerebras_llm():
    if not CEREBRAS_API_KEY:
        raise RuntimeError("CEREBRAS_API_KEY is not set in the environment.")
    return OpenAI(
        api_key=CEREBRAS_API_KEY,
        base_url=CEREBRAS_BASE_URL,
    )


def load_openai_llm():
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not set in the environment.")
    return ChatOpenAI(model=OPENAI_MODEL_NAME)


def read_srt(path: Path) -> str:
    lines: List[str] = []
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if not line or line.isdigit() or "-->" in line:
                continue
            lines.append(line)
    return "\n".join(lines)


def read_vtt(path: Path) -> str:
    lines: List[str] = []
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if (
                not line
                or line.startswith("WEBVTT")
                or line.startswith("NOTE")
                or line.startswith("STYLE")
                or line.startswith("REGION")
                or "-->" in line
                or line.isdigit()
            ):
                continue
            lines.append(line)
    return "\n".join(lines)


def read_txt(path: Path) -> str:
    with path.open("r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_transcript(path: Path) -> str:
    if path.suffix.lower() == ".srt":
        return read_srt(path)
    if path.suffix.lower() == ".vtt":
        return read_vtt(path)
    return read_txt(path)


def chunk_text(text: str, length: int = CHUNK_CHAR_LENGTH, overlap: int = CHUNK_CHAR_OVERLAP) -> List[str]:
    text = text.strip()
    if not text:
        return []

    chunks: List[str] = []
    start = 0
    n = len(text)

    while start < n:
        end = min(start + length, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)

    return chunks


def find_transcript_for_day(day_name: str) -> Path | None:
    """Find the processed transcript file for a given day name."""
    # First check processed transcripts
    for ext in [".txt"]:
        for p in PROCESSED_TRANSCRIPTS_DIR.glob(f"*{day_name}*{ext}"):
            return p

    # If no processed transcript, look for raw transcript
    for ext in [".srt", ".vtt"]:
        for p in TRANSCRIPTS_DIR.glob(f"*{day_name}*{ext}"):
            return p

    return None


# ---------------------------------------------------------------------------
# LLM Call Helpers
# ---------------------------------------------------------------------------

def call_cerebras(llm, system_prompt: str, user_prompt: str) -> str:
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    resp = llm.chat.completions.create(
        model=CEREBRAS_MODEL_NAME,
        messages=messages,
    )
    return resp.choices[0].message.content


def call_openai(llm, system_prompt: str, user_prompt: str) -> str:
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    resp = llm.invoke(messages)
    return resp.content if hasattr(resp, "content") else str(resp)


def summarize_transcript_file(path: Path, cerebras_llm, openai_llm) -> str:
    """Create a topic-based daily summary for a single transcript file."""
    raw = load_transcript(path)
    chunks = chunk_text(raw)
    if not chunks:
        return ""

    partial_summaries: List[str] = []

    for i, chunk in enumerate(chunks, start=1):
        print(f"    - [Cerebras] Summarizing chunk {i}/{len(chunks)} for {path.name}...")
        user_prompt = DAILY_USER_PROMPT_TEMPLATE.format(transcript_chunk=chunk)
        chunk_summary = call_cerebras(cerebras_llm, DAILY_SYSTEM_PROMPT, user_prompt)
        partial_summaries.append(chunk_summary.strip())

    if len(partial_summaries) == 1:
        return partial_summaries[0]

    merged_text = "\n\n---\n\n".join(partial_summaries)
    print(f"    - [OpenAI] Creating final daily summary for {path.name}...")
    user_prompt = DAILY_USER_PROMPT_TEMPLATE.format(transcript_chunk=merged_text)
    final_summary = call_openai(openai_llm, DAILY_SYSTEM_PROMPT, user_prompt)
    return final_summary.strip()


def create_master_summary(openai_llm, daily_summary_paths: List[Path]) -> str:
    """Create a weekly master summary from multiple daily summary files."""
    texts: List[str] = []
    for p in daily_summary_paths:
        with p.open("r", encoding="utf-8", errors="ignore") as f:
            texts.append(f"=== {p.name} ===\n{f.read().strip()}")

    combined = "\n\n\n".join(texts)
    print(f"    📊 [OpenAI] Master summary input: {len(combined):,} chars (~{len(combined)//4:,} tokens)")

    user_prompt = MASTER_USER_PROMPT_TEMPLATE.format(daily_summaries_text=combined)
    master_summary = call_openai(openai_llm, MASTER_SYSTEM_PROMPT, user_prompt)
    return master_summary.strip()


def generate_opening_paragraph(openai_llm, report_content: str) -> str:
    user_prompt = OPENING_USER_PROMPT_TEMPLATE.format(report_content=report_content[:3000])
    return call_openai(openai_llm, OPENING_SYSTEM_PROMPT, user_prompt).strip()


def generate_closing_paragraph(openai_llm, report_content: str) -> str:
    sample_content = report_content[:2000] + "\n...\n" + report_content[-1000:]
    user_prompt = CLOSING_USER_PROMPT_TEMPLATE.format(report_content=sample_content)
    return call_openai(openai_llm, CLOSING_SYSTEM_PROMPT, user_prompt).strip()


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------

def create_word_document(content: str, opening: str, closing: str, output_path: Path) -> None:
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)

    doc.add_paragraph()
    opening_para = doc.add_paragraph(opening)
    opening_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.isupper() or (line[0].isupper() and not line.startswith('•') and not line.startswith('-')):
            if len(doc.paragraphs) > 3:
                doc.add_paragraph()

            heading = doc.add_paragraph(line)
            heading.style = 'Heading 1'
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(12)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            bullet_text = line.lstrip('•-*').strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(6)

        else:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    closing_para = doc.add_paragraph(closing)
    closing_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    signature = doc.add_paragraph("Sincerely,")
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    ensure_dirs()

    print("=" * 70)
    print(" Resume Summarization Script")
    print("=" * 70)
    print(f"\nDays to process: {DAYS_TO_PROCESS if DAYS_TO_PROCESS else 'None (skipping to master summary)'}")
    print(f"Skip daily summaries: {SKIP_DAILY_SUMMARIES}")
    print()

    # Load LLMs
    print("Loading LLMs...")
    print(f"  → Cerebras ({CEREBRAS_MODEL_NAME}) for fast chunk processing")
    print(f"  → OpenAI ({OPENAI_MODEL_NAME}) for final synthesis")
    cerebras_llm = load_cerebras_llm()
    openai_llm = load_openai_llm()
    print()

    # Step 1: Process remaining daily summaries (if any)
    if not SKIP_DAILY_SUMMARIES and DAYS_TO_PROCESS:
        print("Step 1: Creating daily summaries for remaining days\n")

        for day_name in DAYS_TO_PROCESS:
            print(f"▶ Looking for {day_name} transcript...")

            transcript_path = find_transcript_for_day(day_name)
            if not transcript_path:
                print(f"   ⚠️  No transcript found for {day_name}, skipping.")
                continue

            print(f"   Found: {transcript_path.name}")

            try:
                summary_text = summarize_transcript_file(transcript_path, cerebras_llm, openai_llm)
                if not summary_text:
                    print(f"   ⚠️  Empty summary for {day_name}, skipping.")
                    continue

                # Use the stem to create consistent naming
                out_name = f"{transcript_path.stem}_summary.txt"
                out_path = DAILY_SUMMARIES_DIR / out_name
                with out_path.open("w", encoding="utf-8") as f:
                    f.write(summary_text)
                print(f"   ✅ Saved daily summary -> {out_path}")
            except Exception as e:
                print(f"   ❌ Error summarizing {day_name}: {e}")
            print()
    else:
        print("Step 1: Skipping daily summaries (already complete or skip flag set)\n")

    # Step 2: Gather all daily summaries
    print("Step 2: Gathering all daily summaries\n")
    daily_summary_paths = sorted(DAILY_SUMMARIES_DIR.glob("*_summary.txt"))

    if not daily_summary_paths:
        print("❌ No daily summaries found. Cannot create master summary.")
        return

    print(f"   Found {len(daily_summary_paths)} daily summaries:")
    for p in daily_summary_paths:
        print(f"      - {p.name}")
    print()

    # Step 3: Create master summary
    print("Step 3: Creating weekly master summary from all daily summaries\n")
    try:
        master_text = create_master_summary(openai_llm, daily_summary_paths)
        master_out_path = MASTER_SUMMARY_DIR / "master_summary.txt"
        with master_out_path.open("w", encoding="utf-8") as f:
            f.write(master_text)
        print(f"✅ Master summary saved -> {master_out_path}")
    except Exception as e:
        print(f"❌ Error creating master summary: {e}")
        return

    # Step 4: Generate opening paragraph
    print("\nStep 4: Generating opening paragraph\n")
    try:
        opening_text = generate_opening_paragraph(openai_llm, master_text)
        print(f"✅ Opening paragraph generated")
    except Exception as e:
        print(f"❌ Error generating opening: {e}")
        opening_text = ""

    # Step 5: Generate closing paragraph
    print("\nStep 5: Generating closing paragraph\n")
    try:
        closing_text = generate_closing_paragraph(openai_llm, master_text)
        print(f"✅ Closing paragraph generated")
    except Exception as e:
        print(f"❌ Error generating closing: {e}")
        closing_text = ""

    # Step 6: Create Word document
    print("\nStep 6: Creating Word document\n")
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = OUTPUT_DIR / f"Weekly_Consulting_Summary_{timestamp}.docx"
        create_word_document(master_text, opening_text, closing_text, docx_path)
        print(f"✅ Word document created -> {docx_path}")
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")

    print("\n" + "=" * 70)
    print("All done! Your weekly consulting summary is ready.")
    print("=" * 70)


if __name__ == "__main__":
    main()
